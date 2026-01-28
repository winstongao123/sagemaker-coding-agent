"""
Document Tools - Word, Excel, Markdown generation
"""
import os

# Lazy import
Tool = None


def _get_tool_class():
    global Tool
    if Tool is None:
        from core.tools import Tool as T
        Tool = T
    return Tool


def create_word_doc(args: dict, ctx) -> str:
    """Create Word document."""
    filepath = args["filepath"]
    content = args["content"]
    title = args.get("title", "")

    # Handle relative paths
    if not os.path.isabs(filepath):
        filepath = os.path.join(ctx.working_dir, filepath)

    # Ensure .docx extension
    if not filepath.endswith(".docx"):
        filepath += ".docx"

    try:
        from docx import Document

        doc = Document()

        if title:
            doc.add_heading(title, 0)

        # Split content by double newlines for paragraphs
        for para in content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Create directory if needed
        dir_path = os.path.dirname(filepath)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)

        doc.save(filepath)
        return f"Created Word document: {filepath}"

    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"Error creating Word document: {e}"


def create_excel(args: dict, ctx) -> str:
    """Create Excel spreadsheet."""
    filepath = args["filepath"]
    data = args["data"]
    sheet_name = args.get("sheet_name", "Sheet1")

    # Handle relative paths
    if not os.path.isabs(filepath):
        filepath = os.path.join(ctx.working_dir, filepath)

    # Ensure .xlsx extension
    if not filepath.endswith(".xlsx"):
        filepath += ".xlsx"

    try:
        import pandas as pd

        df = pd.DataFrame(data)

        # Create directory if needed
        dir_path = os.path.dirname(filepath)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)

        df.to_excel(filepath, sheet_name=sheet_name, index=False)
        return f"Created Excel file: {filepath} ({len(df)} rows)"

    except ImportError:
        return "Error: pandas/openpyxl not installed. Run: pip install pandas openpyxl"
    except Exception as e:
        return f"Error creating Excel file: {e}"


def create_markdown(args: dict, ctx) -> str:
    """Create Markdown file."""
    filepath = args["filepath"]
    content = args["content"]

    # Handle relative paths
    if not os.path.isabs(filepath):
        filepath = os.path.join(ctx.working_dir, filepath)

    # Ensure .md extension
    if not filepath.endswith(".md"):
        filepath += ".md"

    try:
        # Create directory if needed
        dir_path = os.path.dirname(filepath)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        return f"Created Markdown file: {filepath}"

    except Exception as e:
        return f"Error creating Markdown file: {e}"


# Tool definitions
def get_tools():
    """Get all document tools."""
    Tool = _get_tool_class()

    CREATE_WORD = Tool(
        name="create_word",
        description="Create Word document (.docx). Content is split into paragraphs by double newlines.",
        parameters={
            "type": "object",
            "properties": {
                "filepath": {
                    "type": "string",
                    "description": "Output path (.docx)",
                },
                "content": {
                    "type": "string",
                    "description": "Document content (paragraphs separated by \\n\\n)",
                },
                "title": {
                    "type": "string",
                    "description": "Document title (optional)",
                },
            },
            "required": ["filepath", "content"],
        },
        execute=create_word_doc,
        requires_approval=True,
    )

    CREATE_EXCEL = Tool(
        name="create_excel",
        description="Create Excel spreadsheet (.xlsx). Data should be list of dicts (rows) with consistent keys (columns).",
        parameters={
            "type": "object",
            "properties": {
                "filepath": {
                    "type": "string",
                    "description": "Output path (.xlsx)",
                },
                "data": {
                    "type": "array",
                    "description": "Data as list of dicts, e.g., [{'name': 'Alice', 'age': 30}, {'name': 'Bob', 'age': 25}]",
                    "items": {"type": "object"},
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Sheet name (default: Sheet1)",
                },
            },
            "required": ["filepath", "data"],
        },
        execute=create_excel,
        requires_approval=True,
    )

    CREATE_MARKDOWN = Tool(
        name="create_markdown",
        description="Create Markdown file (.md).",
        parameters={
            "type": "object",
            "properties": {
                "filepath": {
                    "type": "string",
                    "description": "Output path (.md)",
                },
                "content": {
                    "type": "string",
                    "description": "Markdown content",
                },
            },
            "required": ["filepath", "content"],
        },
        execute=create_markdown,
        requires_approval=True,
    )

    return CREATE_WORD, CREATE_EXCEL, CREATE_MARKDOWN


# Export tools
CREATE_WORD, CREATE_EXCEL, CREATE_MARKDOWN = get_tools()
